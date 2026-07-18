"""将 TipTap / ProseMirror JSON 转为 Markdown 或写入 python-docx 文档。"""

from __future__ import annotations

import re
import tempfile
from contextlib import contextmanager
from pathlib import Path
from typing import Any, Iterator
from uuid import UUID

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from sqlalchemy.orm import Session

from app.services.markdown_to_tiptap import _parse_inline_bold
from app.services.publication.publication_styles import DOC_BODY_FONT, FIRST_LINE_INDENT_PT
from app.services.storage_policy import local_business_storage_allowed

BLACK = RGBColor(0, 0, 0)
CODE_SHADING = "F5F5F5"
HEADING_PT = {1: 22, 2: 18, 3: 16, 4: 14, 5: 13, 6: 12}
BODY_PT = 12
CODE_PT = 10


def _inline_to_markdown(nodes: list[dict[str, Any]] | None) -> str:
    if not nodes:
        return ""
    parts: list[str] = []
    for node in nodes:
        t = node.get("type")
        if t == "text":
            text = str(node.get("text") or "")
            marks = node.get("marks") or []
            bold = any(m.get("type") == "bold" for m in marks)
            italic = any(m.get("type") == "italic" for m in marks)
            code = any(m.get("type") == "code" for m in marks)
            if code:
                text = f"`{text}`"
            if bold:
                text = f"**{text}**"
            if italic:
                text = f"*{text}*"
            parts.append(text)
        elif t == "hardBreak":
            parts.append("\n")
        elif t == "mathInline":
            latex = str((node.get("attrs") or {}).get("latex") or "")
            parts.append(f"${latex}$")
        elif t == "citation":
            parts.append(str((node.get("attrs") or {}).get("renderedText") or "（引用）"))
    return "".join(parts)


def _bullet_indent(depth: int) -> str:
    return "  " * max(0, depth)


def _list_item_body(item: dict[str, Any], depth: int) -> str:
    inner = item.get("content") or []
    parts: list[str] = []
    for child in inner:
        if not isinstance(child, dict):
            continue
        ct = child.get("type")
        if ct in ("bulletList", "orderedList"):
            parts.append(_block_to_markdown(child, depth + 1))
        else:
            parts.append(_block_to_markdown(child, depth))
    return "\n\n".join(p for p in parts if p)


def _block_to_markdown(node: dict[str, Any], depth: int = 0) -> str:
    t = node.get("type")
    if t == "paragraph":
        attrs = node.get("attrs") or {}
        pid = str(attrs.get("paragraphId") or "").strip()
        body = _inline_to_markdown(node.get("content"))
        if attrs.get("firstLineIndent") and body.strip() and not body.startswith("\u3000\u3000"):
            body = "\u3000\u3000" + body.lstrip(" \t\u3000")
        return f"<!-- pid:{pid} -->\n{body}" if pid else body
    if t == "heading":
        level = int((node.get("attrs") or {}).get("level") or 1)
        level = max(1, min(6, level))
        return f"{'#' * level} {_inline_to_markdown(node.get('content'))}"
    if t == "blockquote":
        inner = _blocks_to_markdown(node.get("content") or [], depth)
        if not inner:
            return "> "
        return "\n".join(f"> {line}" for line in inner.split("\n"))
    if t == "codeBlock":
        raw = _inline_to_markdown(node.get("content"))
        lang = str((node.get("attrs") or {}).get("language") or "").strip()
        head = f"```{lang}\n" if lang else "```\n"
        return f"{head}{raw}\n```"
    if t == "horizontalRule":
        return "---"
    if t == "bulletList":
        lines: list[str] = []
        for item in node.get("content") or []:
            if not isinstance(item, dict) or item.get("type") != "listItem":
                continue
            body = _list_item_body(item, depth)
            pad = _bullet_indent(depth)
            first = True
            for line in body.split("\n"):
                if first:
                    lines.append(f"{pad}- {line}")
                    first = False
                else:
                    lines.append(f"{pad}  {line}")
        return "\n".join(lines)
    if t == "orderedList":
        start = (node.get("attrs") or {}).get("start")
        idx = int(start) if isinstance(start, int) and start >= 1 else 1
        lines_o: list[str] = []
        for item in node.get("content") or []:
            if not isinstance(item, dict) or item.get("type") != "listItem":
                continue
            body = _list_item_body(item, depth)
            pad = _bullet_indent(depth)
            first = True
            for line in body.split("\n"):
                if first:
                    lines_o.append(f"{pad}{idx}. {line}")
                    first = False
                else:
                    lines_o.append(f"{pad}   {line}")
            idx += 1
        return "\n".join(lines_o)
    if t == "table":
        rows_in = [
            r for r in (node.get("content") or []) if isinstance(r, dict) and r.get("type") == "tableRow"
        ]
        if not rows_in:
            return ""
        md_rows: list[str] = []
        sep_added = False
        for ri, row in enumerate(rows_in):
            cells = [
                c
                for c in (row.get("content") or [])
                if isinstance(c, dict) and c.get("type") in ("tableCell", "tableHeader")
            ]
            texts = []
            for cell in cells:
                parts: list[str] = []
                for sub in cell.get("content") or []:
                    if isinstance(sub, dict) and sub.get("type") == "paragraph":
                        parts.append(_inline_to_markdown(sub.get("content")))
                texts.append(" ".join(p for p in parts if p).strip())
            md_rows.append("| " + " | ".join(texts) + " |")
            if ri == 0 and not sep_added:
                md_rows.append("| " + " | ".join(["---"] * max(1, len(texts))) + " |")
                sep_added = True
        return "\n".join(md_rows)
    if t in ("diagramBlock", "mermaidBlock"):
        attrs = node.get("attrs") or {}
        engine = str(attrs.get("engine") or "graphviz")
        code = str(attrs.get("code") or "").strip()
        lang = "plantuml" if engine == "plantuml" else "dot"
        return f"```{lang}\n{code}\n```"
    if t == "figureBlock":
        attrs = node.get("attrs") or {}
        raw = str(attrs.get("rawAnnotation") or attrs.get("caption") or "").strip()
        ftype = str(attrs.get("figureType") or "figure").lower()
        if ftype == "screenshot":
            tag = "SCREENSHOT"
        elif ftype == "flowchart":
            tag = "FLOWCHART"
        elif ftype == "chart":
            tag = "CHART"
        else:
            tag = "DIAGRAM"
        return f"[{tag}: {raw}]" if raw else f"[{tag}: ]"
    if t == "mathBlock":
        attrs = node.get("attrs") or {}
        latex = str(attrs.get("latex") or "").strip()
        body = f"$$\n{latex}\n$$"
        if attrs.get("numbered") and str(attrs.get("equationNumber") or "").strip():
            body += f"\n<!-- eq-number:{attrs.get('equationNumber')} -->"
        if str(attrs.get("label") or "").strip():
            body += f"\n<!-- eq-label:{attrs.get('label')} -->"
        return body
    if t == "mathInline":
        latex = str((node.get("attrs") or {}).get("latex") or "")
        return f"${latex}$"
    if t == "listItem":
        return _list_item_body(node, depth)
    return ""


def _blocks_to_markdown(nodes: list[dict[str, Any]], depth: int = 0) -> str:
    chunks: list[str] = []
    for node in nodes:
        if not isinstance(node, dict):
            continue
        s = _block_to_markdown(node, depth)
        if s:
            chunks.append(s)
    return "\n\n".join(chunks)


def tiptap_json_to_markdown(doc: dict[str, Any] | None) -> str:
    if not doc or doc.get("type") != "doc":
        return ""
    return _blocks_to_markdown(doc.get("content") or [])


def chapter_content_to_markdown(content: dict[str, Any] | None) -> str:
    if not content or not isinstance(content, dict):
        return ""
    tx = content.get("text")
    text_md = tx.strip("\n\r ") if isinstance(tx, str) else ""
    tj = content.get("tiptap_json")
    if isinstance(tj, dict):
        tj_md = tiptap_json_to_markdown(tj).strip("\n\r ")
        if text_md and "|" in text_md and '"table"' not in str(tj):
            return text_md
        if tj_md:
            return tj_md
    if text_md:
        return text_md
    return ""


# --- DOCX ---


def _apply_run_font(run, name: str) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)


def _style_run(run, *, size_pt: int = BODY_PT, bold: bool = False, italic: bool = False, mono: bool = False) -> None:
    run.font.color.rgb = BLACK
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    _apply_run_font(run, "Consolas" if mono else DOC_BODY_FONT)


def _append_omml_to_paragraph(paragraph, omml_xml: str) -> None:
    """向段落追加 inline OMML（m:oMath 为 w:p 的直接子元素）。"""
    if not omml_xml:
        return
    paragraph._p.append(parse_xml(omml_xml))


def _append_omml_block_to_document(doc: Document, omml_xml: str, *, numbered: bool = False, number: str = "") -> None:
    from app.services.latex_omml import omml_block_xml, omml_numbered_table_xml

    if numbered and number.strip():
        tbl_xml = omml_numbered_table_xml(omml_xml, number.strip())
        doc.element.body.append(parse_xml(tbl_xml))
        return
    p = doc.add_paragraph()
    block = omml_block_xml(omml_xml)
    p._p.append(parse_xml(block))


def _docx_inline_math(paragraph, latex: str) -> None:
    from app.services.latex_omml import latex_to_omml

    result = latex_to_omml(latex)
    if result.get("status") == "ok" and result.get("omml"):
        try:
            _append_omml_to_paragraph(paragraph, str(result["omml"]))
            return
        except Exception:
            pass
    run = paragraph.add_run(f"${latex}$")
    _style_run(run, italic=True)


def _add_inline_to_paragraph(paragraph, nodes: list[dict[str, Any]] | None, *, size_pt: int = BODY_PT) -> None:
    if not nodes:
        return
    for node in nodes:
        t = node.get("type")
        if t == "mathInline":
            _docx_inline_math(paragraph, str((node.get("attrs") or {}).get("latex") or ""))
            continue
        if t == "citation":
            run = paragraph.add_run(str((node.get("attrs") or {}).get("renderedText") or "（引用）"))
            _style_run(run, size_pt=size_pt)
            continue
        if t == "text":
            run = paragraph.add_run(str(node.get("text") or ""))
            bold = italic = mono = False
            for m in node.get("marks") or []:
                mt = m.get("type")
                if mt == "bold":
                    bold = True
                elif mt == "italic":
                    italic = True
                elif mt == "code":
                    mono = True
            _style_run(run, size_pt=CODE_PT if mono else size_pt, bold=bold, italic=italic, mono=mono)
        elif t == "hardBreak":
            run = paragraph.add_run("\n")
            _style_run(run, size_pt=size_pt)


def _body_paragraph_format(p, *, indent: bool = True) -> None:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)
    if indent:
        pf.first_line_indent = Pt(FIRST_LINE_INDENT_PT)


def _set_paragraph_shading(paragraph, fill: str = CODE_SHADING) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _add_code_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(12)
    p.paragraph_format.right_indent = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    _set_paragraph_shading(p)
    run = p.add_run(text)
    _style_run(run, size_pt=CODE_PT, mono=True)


def _add_paragraph_bullet(doc: Document, content: list[dict[str, Any]] | None) -> None:
    try:
        p = doc.add_paragraph(style="List Bullet")
    except Exception:
        p = doc.add_paragraph()
        r = p.add_run("• ")
        _style_run(r)
    _body_paragraph_format(p)
    _add_inline_to_paragraph(p, content)


def _add_paragraph_numbered(doc: Document, content: list[dict[str, Any]] | None) -> None:
    try:
        p = doc.add_paragraph(style="List Number")
    except Exception:
        p = doc.add_paragraph()
    _body_paragraph_format(p)
    _add_inline_to_paragraph(p, content)


def _table_cell_text(cell: dict[str, Any]) -> str:
    parts: list[str] = []
    for sub in cell.get("content") or []:
        if isinstance(sub, dict) and sub.get("type") == "paragraph":
            parts.append(_inline_to_markdown(sub.get("content")))
    return " ".join(p for p in parts if p).strip()


def _expand_bold_in_inline_nodes(nodes: list[dict[str, Any]]) -> list[dict[str, Any]]:
    out: list[dict[str, Any]] = []
    for node in nodes:
        if not isinstance(node, dict) or node.get("type") != "text":
            out.append(node)
            continue
        marks = node.get("marks") or []
        text = str(node.get("text") or "")
        if not marks and "**" in text:
            out.extend(_parse_inline_bold(text))
        else:
            out.append(node)
    return out


def _table_cell_inline_nodes(cell: dict[str, Any]) -> list[dict[str, Any]]:
    """提取单元格内联节点；若仅有字面量 ** 则解析为 bold mark。"""
    nodes: list[dict[str, Any]] = []
    for sub in cell.get("content") or []:
        if isinstance(sub, dict) and sub.get("type") == "paragraph":
            content = sub.get("content") or []
            if isinstance(content, list):
                nodes.extend(content)
    if nodes:
        return _expand_bold_in_inline_nodes(nodes)
    text = _table_cell_text(cell)
    if not text:
        return []
    if "**" in text:
        return _parse_inline_bold(text)
    return [{"type": "text", "text": text}]


_FIGURE_URL_RE = re.compile(
    r"/static/figures/([0-9a-f-]+)/(\d+)/([0-9a-f-]+)/figure\.(\w+)",
    re.I,
)
_RASTER_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".webp"}


def _figure_path_rank(path: Path) -> int:
    ext = path.suffix.lower()
    if ext == ".png":
        return 0
    if ext in {".jpg", ".jpeg", ".webp"}:
        return 1
    if ext == ".gif":
        return 2
    if ext == ".svg":
        return 3
    return 4


def _strip_figure_url(raw: str) -> str:
    return str(raw or "").strip().split("?", 1)[0].split("#", 1)[0].strip()


def _parse_figure_ids(attrs: dict[str, Any]) -> tuple[str, str | None, int | None]:
    figure_id = str(attrs.get("figureId") or attrs.get("figure_id") or "").strip()
    book_id = str(attrs.get("book_id") or attrs.get("bookId") or "").strip() or None
    chapter_raw = attrs.get("chapter_index", attrs.get("chapterIndex"))
    chapter_index: int | None = (
        int(chapter_raw) if chapter_raw is not None and str(chapter_raw) != "" else None
    )
    for key in ("fileUrl", "file_url", "svgUrl", "svg_url", "file_path"):
        raw = _strip_figure_url(str(attrs.get(key) or ""))
        match = _FIGURE_URL_RE.search(raw)
        if match:
            book_id = book_id or match.group(1)
            if not figure_id:
                figure_id = match.group(3)
            if chapter_index is None:
                chapter_index = int(match.group(2))
            break
    return figure_id, book_id, chapter_index


def merge_figure_export_attrs(
    block_attrs: dict[str, Any],
    node_attrs: dict[str, Any] | None,
) -> dict[str, Any]:
    """合并导出用 figure 属性，避免 TipTap 中空 URL 覆盖数据库/AST 中的有效路径。"""
    base = {k: v for k, v in block_attrs.items() if k != "tiptap_node"}
    merged = {**base, **(node_attrs or {})}
    for key in ("fileUrl", "file_url", "svgUrl", "svg_url", "file_path"):
        block_val = str(base.get(key) or "").strip()
        node_val = str((node_attrs or {}).get(key) or "").strip()
        if block_val and not node_val:
            merged[key] = base[key]
    return merged


def _canonical_figure_paths(attrs: dict[str, Any]) -> list[Path]:
    from app.services.figures.export_assets import find_figure_asset_on_disk
    from app.services.figures.storage.manager import figure_storage

    figure_id, book_id, chapter_index = _parse_figure_ids(attrs)
    file_path = str(attrs.get("file_path") or "").strip()
    if file_path:
        path = Path(file_path)
        if path.is_file():
            return [path]
    if not figure_id or not book_id:
        return []
    try:
        fid = UUID(figure_id)
        book_uuid = UUID(book_id)
    except ValueError:
        return []
    if not local_business_storage_allowed():
        return []
    paths: list[Path] = []
    if chapter_index is not None:
        paths.extend(
            [
                figure_storage.png_path(book_uuid, chapter_index, fid),
                figure_storage.svg_path(book_uuid, chapter_index, fid),
            ]
        )
    paths.extend(find_figure_asset_on_disk(book_uuid, fid))
    return [p for p in paths if p.is_file()]


def _resolve_figure_local_path(attrs: dict[str, Any], db: Session | None = None) -> Path | None:
    if db is not None:
        from app.services.assets.asset_resolver import AssetResolver

        resolved = AssetResolver(db).resolve_figure_local_path_from_attrs(attrs)
        if resolved and resolved.is_file():
            return resolved

    from app.services.figures.storage.manager import figure_storage

    candidates: list[Path] = []
    seen: set[str] = set()

    def _add(path: Path | None) -> None:
        if not path or not path.is_file():
            return
        key = str(path.resolve())
        if key in seen:
            return
        seen.add(key)
        candidates.append(path)

    for key in ("fileUrl", "file_url", "svgUrl", "svg_url", "file_path"):
        raw = _strip_figure_url(str(attrs.get(key) or ""))
        if raw:
            if key == "file_path":
                _add(Path(raw) if Path(raw).is_file() else None)
            else:
                _add(figure_storage.resolve_local_path(raw))

    for path in _canonical_figure_paths(attrs):
        _add(path)

    if not candidates:
        return None
    return min(candidates, key=_figure_path_rank)


@contextmanager
def _materialize_figure_local_path(
    attrs: dict[str, Any],
    db: Session | None = None,
) -> Iterator[Path | None]:
    if db is not None:
        from app.services.assets.asset_resolver import AssetResolver

        with AssetResolver(db).materialize_figure_local_path_from_attrs(attrs) as resolved:
            if resolved and resolved.is_file():
                yield resolved
                return

    local = _resolve_figure_local_path(attrs, db=None)
    yield local if local and local.is_file() else None


def _new_temp_path(*, suffix: str) -> Path:
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        return Path(tmp.name)


def _figure_raster_for_export_owned(local: Path) -> tuple[Path | None, bool]:
    """Return a DOCX/PDF-compatible raster and whether the caller owns it."""
    ext = local.suffix.lower()
    if ext in _RASTER_EXTS:
        return (local if local.is_file() else None), False

    if ext == ".svg":
        png_candidates: list[Path] = []
        if local.name == "figure.svg":
            png_candidates.append(local.with_name("figure.png"))
        png_candidates.append(local.with_suffix(".png"))
        for png_candidate in png_candidates:
            if png_candidate.is_file():
                return _figure_raster_for_export_owned(png_candidate)
        from app.services.figures.render.svg.export_png import export_png_from_svg

        if local.name == "figure.svg":
            canonical_png = local.with_name("figure.png")
            if local_business_storage_allowed() and export_png_from_svg(local, canonical_png) and canonical_png.is_file():
                return canonical_png, False

        cache_path = _new_temp_path(suffix=".svg-raster.png")
        if export_png_from_svg(local, cache_path) and cache_path.is_file():
            return cache_path, True
        cache_path.unlink(missing_ok=True)
    return None, False


@contextmanager
def _materialize_figure_raster_for_export(local: Path) -> Iterator[Path | None]:
    raster, owned = _figure_raster_for_export_owned(local)
    try:
        yield raster
    finally:
        if owned and raster:
            raster.unlink(missing_ok=True)


def _figure_raster_for_export(local: Path) -> Path | None:
    """DOCX/PDF 仅支持位图；SVG 尝试转 PNG 或使用同目录 figure.png。"""
    return _figure_raster_for_export_owned(local)[0]


def _docx_figure_size(local: Path) -> tuple[float, float | None]:
    """Return width/height in inches while preserving the image aspect ratio."""
    max_width = 5.5
    max_height = 5.8
    square_width = 4.2
    try:
        from PIL import Image

        with Image.open(local) as img:
            px_w, px_h = img.size
    except Exception:
        return 5.0, None

    if px_w <= 0 or px_h <= 0:
        return 5.0, None

    aspect = px_w / px_h
    if 0.85 <= aspect <= 1.18:
        width = square_width
    else:
        width = max_width
    height = width / aspect
    if height > max_height:
        height = max_height
        width = min(max_width, height * aspect)
    width = max(2.4, min(max_width, width))
    height = max(1.2, min(max_height, height))
    return width, height


def docx_figure_image_only(doc: Document, node: dict[str, Any], db: Session | None = None) -> bool:
    """仅插入图片，不写「图 1-1」「图解：」等（题注由 AST figure_caption 负责）。"""
    attrs = node.get("attrs") or {}
    with _materialize_figure_local_path(attrs, db=db) as local:
        if not local:
            return False
        with _materialize_figure_raster_for_export(local) as raster:
            if not raster:
                return False
            pic_p = doc.add_paragraph()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            width, height = _docx_figure_size(raster)
            kwargs: dict[str, Any] = {"width": Inches(width)}
            if height is not None:
                kwargs["height"] = Inches(height)
            try:
                pic_p.add_run().add_picture(str(raster), **kwargs)
            except Exception:
                return False
            return True


def _docx_block(doc: Document, node: dict[str, Any]) -> None:
    t = node.get("type")
    if t == "paragraph":
        attrs = node.get("attrs") or {}
        center = attrs.get("textAlign") == "center"
        p = doc.add_paragraph()
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _body_paragraph_format(p, indent=not center)
        _add_inline_to_paragraph(p, node.get("content"))
        return
    if t == "heading":
        level = int((node.get("attrs") or {}).get("level") or 1)
        level = max(1, min(6, level))
        try:
            p = doc.add_paragraph(style=f"Heading {level}")
        except Exception:
            try:
                p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
            except Exception:
                p = doc.add_paragraph()
        _add_inline_to_paragraph(p, node.get("content"), size_pt=HEADING_PT.get(level, 14))
        for run in p.runs:
            run.bold = True
        return
    if t == "table":
        rows_in = [r for r in (node.get("content") or []) if isinstance(r, dict) and r.get("type") == "tableRow"]
        if not rows_in:
            return
        max_cols = 0
        for row in rows_in:
            cells = [c for c in (row.get("content") or []) if isinstance(c, dict)]
            max_cols = max(max_cols, len(cells))
        max_cols = max(1, max_cols)
        table = doc.add_table(rows=len(rows_in), cols=max_cols)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        for ri, row in enumerate(rows_in):
            cells = [c for c in (row.get("content") or []) if isinstance(c, dict)]
            for ci in range(max_cols):
                cell = cells[ci] if ci < len(cells) else {}
                cell_p = table.rows[ri].cells[ci].paragraphs[0]
                cell_p.clear()
                inline_nodes = _table_cell_inline_nodes(cell) if cell else []
                if inline_nodes:
                    _add_inline_to_paragraph(cell_p, inline_nodes, size_pt=BODY_PT)
        doc.add_paragraph("")
        return
    if t == "bulletList":
        for item in node.get("content") or []:
            if not isinstance(item, dict) or item.get("type") != "listItem":
                continue
            for sub in item.get("content") or []:
                if isinstance(sub, dict) and sub.get("type") == "paragraph":
                    _add_paragraph_bullet(doc, sub.get("content"))
                elif isinstance(sub, dict):
                    _docx_block(doc, sub)
        return
    if t == "orderedList":
        for item in node.get("content") or []:
            if not isinstance(item, dict) or item.get("type") != "listItem":
                continue
            for sub in item.get("content") or []:
                if isinstance(sub, dict) and sub.get("type") == "paragraph":
                    _add_paragraph_numbered(doc, sub.get("content"))
                elif isinstance(sub, dict):
                    _docx_block(doc, sub)
        return
    if t == "blockquote":
        for sub in node.get("content") or []:
            if isinstance(sub, dict) and sub.get("type") == "paragraph":
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Pt(18)
                _body_paragraph_format(p)
                _add_inline_to_paragraph(p, sub.get("content"))
        return
    if t == "codeBlock":
        _add_code_paragraph(doc, _inline_to_markdown(node.get("content")))
        return
    if t in ("diagramBlock", "mermaidBlock"):
        attrs = node.get("attrs") or {}
        code = str(attrs.get("code") or "").strip()
        lang = str(attrs.get("engine") or "mermaid")
        _add_code_paragraph(doc, f"[{lang}]\n{code}")
        return
    if t == "horizontalRule":
        p = doc.add_paragraph("—" * 24)
        for run in p.runs:
            _style_run(run)
        return
    if t == "figureBlock":
        if not docx_figure_image_only(doc, node):
            attrs = node.get("attrs") or {}
            num = str(attrs.get("figureNumber") or "")
            label = f"图{num}" if num else "图"
            p = doc.add_paragraph(f"【{label} 待生成】")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                _style_run(run)
        return
    if t == "mathBlock":
        attrs = node.get("attrs") or {}
        latex = str(attrs.get("latex") or "").strip()
        if not latex:
            return
        from app.services.latex_omml import latex_to_omml

        result = latex_to_omml(latex)
        if result.get("status") == "ok" and result.get("omml"):
            try:
                _append_omml_block_to_document(
                    doc,
                    str(result["omml"]),
                    numbered=bool(attrs.get("numbered")),
                    number=str(attrs.get("equationNumber") or ""),
                )
                return
            except Exception:
                pass
        p = doc.add_paragraph(f"${latex}$")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            _style_run(run, italic=True)
        return


def append_tiptap_to_document(doc: Document, tiptap: dict[str, Any] | None) -> None:
    if not tiptap or tiptap.get("type") != "doc":
        return
    for node in tiptap.get("content") or []:
        if isinstance(node, dict):
            docx_block(doc, node)


docx_block = _docx_block


def append_chapter_content_to_document(doc: Document, content: dict[str, Any] | None) -> None:
    if not content or not isinstance(content, dict):
        return
    tj = content.get("tiptap_json")
    if isinstance(tj, dict):
        append_tiptap_to_document(doc, tj)
        return
    tx = content.get("text")
    if isinstance(tx, str) and tx.strip():
        for para in tx.strip().split("\n\n"):
            doc.add_paragraph(para.replace("\n", " "))
