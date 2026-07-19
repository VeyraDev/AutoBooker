"""Export page numbers: Arabic #1 starts at body (正文), front matter unnumbered."""

from __future__ import annotations

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def _append_page_field(run, *, placeholder: str = "1") -> None:
    """插入 PAGE 域；separate 与 end 之间放占位数字，打开即可见。"""
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE \\* MERGEFORMAT "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(text)
    run._r.append(fld_end)


def _set_section_page_start(section, start: int) -> None:
    sect_pr = section._sectPr
    pg = sect_pr.find(qn("w:pgNumType"))
    if pg is None:
        pg = OxmlElement("w:pgNumType")
        sect_pr.append(pg)
    pg.set(qn("w:start"), str(start))


def _disable_title_page(section) -> None:
    try:
        section.different_first_page_header_footer = False
    except Exception:
        pass
    sect_pr = section._sectPr
    title_pg = sect_pr.find(qn("w:titlePg"))
    if title_pg is not None:
        sect_pr.remove(title_pg)


def clear_section_footer(section) -> None:
    def _clear(footer) -> None:
        footer.is_linked_to_previous = False
        for p in list(footer.paragraphs):
            p_element = p._p
            for child in list(p_element):
                if child.tag.endswith(("r", "fldSimple")):
                    p_element.remove(child)

    _clear(section.footer)
    try:
        _clear(section.even_page_footer)
    except Exception:
        pass


def set_section_footer_page_number(section) -> None:
    _disable_title_page(section)

    def _fill(footer) -> None:
        footer.is_linked_to_previous = False
        # 清空旧段落，重建唯一居中页码段
        for p in list(footer.paragraphs):
            p_element = p._p
            for child in list(p_element):
                if child.tag.endswith(("r", "fldSimple")):
                    p_element.remove(child)
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run()
        _append_page_field(run, placeholder="1")
        run.font.size = Pt(10.5)
        run.font.name = "Times New Roman"
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), "Times New Roman")
        rfonts.set(qn("w:hAnsi"), "Times New Roman")
        rfonts.set(qn("w:eastAsia"), "宋体")

    _fill(section.footer)
    try:
        _fill(section.even_page_footer)
    except Exception:
        pass


def enable_update_fields_on_open(doc: Document) -> None:
    """打开文档时更新域，避免 WPS/Word 页脚一直显示空白。"""
    settings_el = doc.settings.element
    existing = settings_el.find(qn("w:updateFields"))
    if existing is not None:
        settings_el.remove(existing)
    el = OxmlElement("w:updateFields")
    el.set(qn("w:val"), "true")
    settings_el.append(el)


def add_docx_page_numbers(doc: Document) -> None:
    for section in doc.sections:
        set_section_footer_page_number(section)
    enable_update_fields_on_open(doc)


def configure_front_and_body_page_numbers(doc: Document, *, body_section_index: int) -> None:
    """
    - 正文前不排阿拉伯页码
    - 正文第一节起页码从 1 开始
    """
    for i, section in enumerate(doc.sections):
        _disable_title_page(section)
        if i < body_section_index:
            clear_section_footer(section)
        else:
            if i == body_section_index:
                _set_section_page_start(section, 1)
            set_section_footer_page_number(section)
    enable_update_fields_on_open(doc)


def detect_pdf_body_start_page(
    pdf_bytes: bytes,
    *,
    chapter_titles: list[str] | None = None,
    fallback: int = 0,
) -> int:
    """定位正文起始页：先找标记，再跳过目录/前言后匹配章名。"""
    titles = [t.strip() for t in (chapter_titles or []) if (t or "").strip()]
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        n = doc.page_count
        if n <= 0:
            return 0

        texts = [(i, (doc.load_page(i).get_text("text") or "")) for i in range(n)]

        for i, text in texts:
            if "[[BODY_START]]" in text or "⟦BODY⟧" in text:
                return i

        toc_idx = None
        for i, text in texts:
            head = text.strip()[:80]
            if head.startswith("目录") or "\n目录\n" in f"\n{text[:120]}\n":
                toc_idx = i
                break

        search_from = (toc_idx + 1) if toc_idx is not None else 0
        for i in range(search_from, n):
            text = texts[i][1]
            head = text.strip()[:40]
            if head.startswith("前言"):
                continue
            # 仍带「目录」的续页跳过
            if "目录" in text[:30] and any(t in text for t in titles[:1]):
                # 目录续页也含章名，不能算正文
                if text.strip().startswith("目录") or text.count("……") + text.count("…") > 3:
                    continue
            for title in titles:
                if title and title in text:
                    return i

        fb = max(0, min(int(fallback), max(0, n - 1)))
        return fb
    finally:
        doc.close()


def scrub_pdf_body_marker(pdf_bytes: bytes) -> bytes:
    """去掉正文起点探测标记，避免印在页面上。"""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        for page in doc:
            found = False
            for needle in ("[[BODY_START]]", "⟦BODY⟧"):
                for rect in page.search_for(needle):
                    page.add_redact_annot(rect, fill=(1, 1, 1))
                    found = True
            if found:
                page.apply_redactions()
        return doc.tobytes()
    finally:
        doc.close()


def add_pdf_page_numbers(
    pdf_bytes: bytes,
    *,
    body_start_page_index: int = 0,
) -> bytes:
    """仅从正文起始页开始印阿拉伯数字；正文第 1 页为 1。"""
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        n = doc.page_count
        start = max(0, min(int(body_start_page_index), n))
        for index in range(start, n):
            page = doc.load_page(index)
            rect = page.rect
            box = fitz.Rect(
                rect.x0 + 40,
                rect.y1 - 40,
                rect.x1 - 40,
                rect.y1 - 14,
            )
            label = str(index - start + 1)
            inserted = False
            for fontname in ("china-s", "china-ss", "helv"):
                try:
                    rc = page.insert_textbox(
                        box,
                        label,
                        fontsize=10.5,
                        fontname=fontname,
                        color=(0, 0, 0),
                        align=fitz.TEXT_ALIGN_CENTER,
                    )
                    if rc >= 0:
                        inserted = True
                        break
                except Exception:
                    continue
            if not inserted:
                try:
                    page.insert_text(
                        (rect.width / 2 - 6, rect.height - 24),
                        label,
                        fontsize=10.5,
                        fontname="helv",
                        color=(0, 0, 0),
                    )
                except Exception:
                    pass
        return doc.tobytes()
    finally:
        doc.close()
