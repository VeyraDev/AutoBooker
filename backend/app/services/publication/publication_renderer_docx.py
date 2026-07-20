"""Render BookAst to DOCX with publication styles."""

from __future__ import annotations

import io

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from sqlalchemy.orm import Session

from app.services.publication.book_ast import AstBlock, BookAst
from app.services.publication.export_ast import BookExportAst
from app.services.publication.page_formats import PageFormatSpec, get_page_format, get_page_format_from_publication
from app.services.publication.page_numbers import (
    add_docx_page_numbers,
    configure_front_and_body_page_numbers,
)
from app.services.publication.publication_styles import (
    AST_WORD_HEADING_LEVEL,
    BODY_PT,
    BOOK_TITLE_PT,
    CAPTION_PT,
    CHAPTER_TITLE_PT,
    DOC_BODY_FONT,
    DOC_HEADING_FONT,
    DOC_PREFACE_FONT,
    FIFTH_TITLE_PT,
    FIRST_LINE_INDENT_PT,
    FOURTH_TITLE_PT,
    SECTION_TITLE_PT,
    SIXTH_TITLE_PT,
    SUBSECTION_TITLE_PT,
    subsection_word_heading_level,
    type_scale_for_format,
)
from app.services.tiptap_convert import (
    _add_code_paragraph,
    _add_inline_to_paragraph,
    append_tiptap_to_document,
    docx_block,
    docx_figure_image_only,
    merge_figure_export_attrs,
)

BLACK = RGBColor(0, 0, 0)


def _set_run_font(run, *, name: str = DOC_BODY_FONT, size_pt: float | None = None) -> None:
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)


def _style_docx_run(
    run,
    *,
    size_pt: float,
    bold: bool = False,
    font_name: str = DOC_BODY_FONT,
) -> None:
    _set_run_font(run, name=font_name, size_pt=size_pt)
    run.bold = bold
    run.font.color.rgb = BLACK


def _apply_page_geometry(section, spec: PageFormatSpec | None = None) -> None:
    """按所选开本设置成品尺寸与出版页边距（订口/切口/天头/地脚）。"""
    spec = spec or get_page_format()
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Mm(spec.width_mm)
    section.page_height = Mm(spec.height_mm)
    # 镜像页边距模式下：left=内侧(订口)，right=外侧(切口)
    section.left_margin = Mm(spec.margin_inner_mm)
    section.right_margin = Mm(spec.margin_outer_mm)
    section.top_margin = Mm(spec.margin_top_mm)
    section.bottom_margin = Mm(spec.margin_bottom_mm)
    section.gutter = Mm(0)  # 镜像模式已含订口加宽；厚书可另加
    # 页码约在地脚上方 5mm；书眉在天头下方 5mm
    section.footer_distance = Mm(8)
    section.header_distance = Mm(5)
    try:
        section.different_first_page_header_footer = False
    except Exception:
        pass


def _enable_mirror_margins(doc: Document) -> None:
    """布局 → 页边距 → 多页：镜像页边距。"""
    settings_el = doc.settings.element
    if settings_el.find(qn("w:mirrorMargins")) is None:
        settings_el.append(OxmlElement("w:mirrorMargins"))
    # 页码目前居中：不要开奇偶页不同页脚，否则偶数页会丢页码
    try:
        doc.settings.odd_and_even_pages_header_footer = False
    except Exception:
        pass


def _init_docx_publication_fonts(doc: Document, spec: PageFormatSpec | None = None) -> None:
    spec = spec or get_page_format()
    scale = type_scale_for_format(spec)
    normal = doc.styles["Normal"]
    normal.font.name = DOC_BODY_FONT
    normal.font.size = Pt(scale.body_pt)
    normal.font.color.rgb = BLACK
    rpr = normal._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), DOC_BODY_FONT)
    rfonts.set(qn("w:hAnsi"), DOC_BODY_FONT)
    rfonts.set(qn("w:eastAsia"), DOC_BODY_FONT)
    _configure_builtin_heading_styles(doc, scale)
    _enable_mirror_margins(doc)
    for section in doc.sections:
        _apply_page_geometry(section, spec)


def _configure_builtin_heading_styles(doc: Document, scale=None) -> None:
    """为内置 Heading 1–6 套用出版字体（黑体分级字号）。"""
    if scale is None:
        scale = type_scale_for_format()
    heading_specs = {
        1: (scale.chapter_pt, True),
        2: (scale.section_pt, True),
        3: (scale.subsection_pt, True),
        4: (scale.fourth_pt, True),
        5: (scale.fifth_pt, True),
        6: (scale.fifth_pt, False),
    }
    for word_level, (size_pt, bold) in heading_specs.items():
        style_name = f"Heading {word_level}"
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = DOC_HEADING_FONT
        style.font.size = Pt(size_pt)
        style.font.bold = bold
        style.font.color.rgb = BLACK
        rpr = style._element.get_or_add_rPr()
        rfonts = rpr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:ascii"), DOC_HEADING_FONT)
        rfonts.set(qn("w:hAnsi"), DOC_HEADING_FONT)
        rfonts.set(qn("w:eastAsia"), DOC_HEADING_FONT)
        _set_style_outline_level(style, word_level - 1)


def _set_style_outline_level(style, outline_level: int) -> None:
    ppr = style._element.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(outline_level))


def _set_paragraph_outline_level(paragraph, outline_level: int) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        ppr.append(outline)
    outline.set(qn("w:val"), str(outline_level))


def _apply_word_heading_style(doc: Document, paragraph, word_level: int) -> None:
    style_name = f"Heading {word_level}"
    paragraph.style = doc.styles[style_name]
    _set_paragraph_outline_level(paragraph, word_level - 1)


def _add_body(doc: Document, text: str, *, scale=None) -> None:
    scale = scale or type_scale_for_format()
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(scale.first_indent_pt)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = scale.line_spacing
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    _style_docx_run(run, size_pt=scale.body_pt)


def _heading_visual(scale, word_level: int):
    """返回 (size, bold, align, space_before, space_after)。二级起左对齐。"""
    table = {
        1: (scale.chapter_pt, True, WD_ALIGN_PARAGRAPH.CENTER, 0, 12),
        2: (scale.section_pt, True, WD_ALIGN_PARAGRAPH.LEFT, 12, 8),
        3: (scale.subsection_pt, True, WD_ALIGN_PARAGRAPH.LEFT, 10, 6),
        4: (scale.fourth_pt, True, WD_ALIGN_PARAGRAPH.LEFT, 8, 4),
        5: (scale.fifth_pt, True, WD_ALIGN_PARAGRAPH.LEFT, 6, 4),
        6: (scale.fifth_pt, False, WD_ALIGN_PARAGRAPH.LEFT, 6, 4),
    }
    return table.get(word_level, table[3])


def _heading_level(block: AstBlock, fallback: int) -> int:
    node = block.attrs.get("tiptap_node")
    if isinstance(node, dict):
        raw = (node.get("attrs") or {}).get("level")
        try:
            return max(1, min(6, int(raw)))
        except Exception:
            pass
    return max(1, min(6, block.level or fallback))


def _word_heading_level_for_block(block: AstBlock) -> int:
    mapped = AST_WORD_HEADING_LEVEL.get(block.role)
    if mapped is not None:
        return mapped
    if block.role == "subsection_title":
        return subsection_word_heading_level(_heading_level(block, 3))
    return subsection_word_heading_level(_heading_level(block, 2))


def _add_publication_heading(
    doc: Document,
    text: str,
    *,
    word_level: int,
    node: dict | None = None,
    scale=None,
) -> None:
    scale = scale or type_scale_for_format()
    word_level = max(1, min(6, word_level))
    size_pt, bold, align, space_before, space_after = _heading_visual(scale, word_level)
    if isinstance(node, dict) and node.get("content"):
        p = doc.add_paragraph()
        _apply_word_heading_style(doc, p, word_level)
        _add_inline_to_paragraph(p, node.get("content"), size_pt=size_pt)
    else:
        p = doc.add_heading(text, level=word_level)
        _set_paragraph_outline_level(p, word_level - 1)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    for r in p.runs:
        _style_docx_run(r, size_pt=size_pt, bold=bold, font_name=DOC_HEADING_FONT)


def _add_chapter_flyleaf(doc: Document, title: str, summary: str, *, scale) -> None:
    """章过渡页：章名 + 一句导语。"""
    for _ in range(6):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(title)
    _style_docx_run(run, size_pt=scale.chapter_pt, bold=True, font_name=DOC_HEADING_FONT)
    # 隐式 outline：章标题入目录
    _set_paragraph_outline_level(p, 0)
    if summary:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.left_indent = Pt(36)
        s.paragraph_format.right_indent = Pt(36)
        s.paragraph_format.line_spacing = 1.7
        sr = s.add_run(summary)
        _style_docx_run(sr, size_pt=scale.body_pt, font_name=DOC_PREFACE_FONT)


def _add_section_flyleaf(doc: Document, title: str, *, scale) -> None:
    """节分隔页：仅节标题居中。"""
    for _ in range(8):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    _style_docx_run(run, size_pt=scale.section_pt, bold=True, font_name=DOC_HEADING_FONT)
    _set_paragraph_outline_level(p, 1)


def _add_page_break(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _add_section_break_new(doc: Document, spec: PageFormatSpec | None = None):
    """Start a new Word section (for page-number restart)."""
    new_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    _apply_page_geometry(new_sec, spec)
    return new_sec


def _add_cover_image(doc: Document, pub: dict, spec: PageFormatSpec, *, db: Session | None = None) -> None:
    from app.services.publication.cover_background import render_cover_png

    png = render_cover_png(pub, fallback_title=pub.get("title"), db=db)
    section = doc.sections[-1]
    old = (
        section.left_margin,
        section.right_margin,
        section.top_margin,
        section.bottom_margin,
    )
    section.left_margin = Mm(0)
    section.right_margin = Mm(0)
    section.top_margin = Mm(0)
    section.bottom_margin = Mm(0)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(io.BytesIO(png), width=Mm(spec.width_mm), height=Mm(spec.height_mm))
    section.left_margin, section.right_margin, section.top_margin, section.bottom_margin = old


def _add_colophon_page(doc: Document, pub: dict, spec: PageFormatSpec) -> None:
    """规范版权页：信息集中在页面下半部（大众书常见排法）。"""
    from app.services.publication.publication_info import build_colophon_lines

    body_pt = spec.body_pt or BODY_PT

    for _ in range(10):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(6)

    expect_cip_body = False
    for line in build_colophon_lines(pub):
        if line == "图书在版编目（CIP）数据":
            p = doc.add_paragraph()
            r = p.add_run(line)
            _style_docx_run(r, size_pt=body_pt - 0.5, bold=True)
            expect_cip_body = True
            continue
        if expect_cip_body:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(10)
            r = p.add_run(line)
            _style_docx_run(r, size_pt=body_pt - 1)
            expect_cip_body = False
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.35
        run = p.add_run(line)
        _style_docx_run(run, size_pt=body_pt - 0.5)


def _add_cover_section(doc: Document, section, spec: PageFormatSpec, *, db: Session | None = None) -> None:
    pub = section.publication if isinstance(getattr(section, "publication", None), dict) else {}
    if not pub:
        pub = {"title": section.title or "未命名"}
    _add_cover_image(doc, pub, spec, db=db)
    _add_page_break(doc)
    _add_colophon_page(doc, pub, spec)


def _add_toc_leader_paragraph(
    doc: Document,
    title: str,
    page: int | None,
    *,
    level: int,
    spec: PageFormatSpec,
) -> None:
    """目录行：标题+点线+等宽页码字段（与 PDF 同一算法）。"""
    from app.services.publication.toc_format import toc_entry_plain_line

    body_pt = spec.body_pt or BODY_PT
    size = body_pt if level <= 1 else max(9.0, body_pt - 0.5)
    indent_pt = 0.0 if level <= 1 else float(FIRST_LINE_INDENT_PT)
    left, page_field = toc_entry_plain_line(
        title,
        page,
        content_width_pt=float(spec.content_width_pt),
        font_size_pt=float(size),
        indent_pt=indent_pt,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    if indent_pt:
        p.paragraph_format.left_indent = Pt(indent_pt)
    run = p.add_run(f"{left}{page_field}" if page_field else left)
    _style_docx_run(run, size_pt=size, bold=level <= 1)


def _add_toc_section(doc: Document, section, spec: PageFormatSpec) -> None:
    p = doc.add_paragraph("目录")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        _style_docx_run(r, size_pt=CHAPTER_TITLE_PT, bold=True)
    for entry in section.entries:
        _add_toc_leader_paragraph(
            doc,
            entry.title,
            entry.page,
            level=getattr(entry, "level", 1) or 1,
            spec=spec,
        )


def render_export_ast_to_docx(export_ast: BookExportAst, db: Session | None = None) -> bytes:
    pub = {}
    for section in export_ast.sections:
        if section.type == "cover" and isinstance(getattr(section, "publication", None), dict):
            pub = section.publication
            break
    spec = get_page_format_from_publication(pub)
    scale = type_scale_for_format(spec)

    doc = Document()
    _init_docx_publication_fonts(doc, spec)
    has_content = False
    last_was_page_break = False
    body_started = False
    body_section_index = 0

    def mark_content() -> None:
        nonlocal has_content, last_was_page_break
        has_content = True
        last_was_page_break = False

    def add_page_break_once() -> None:
        nonlocal last_was_page_break
        if not has_content or last_was_page_break:
            return
        _add_page_break(doc)
        last_was_page_break = True

    def begin_body_if_needed() -> None:
        nonlocal body_started, body_section_index, last_was_page_break
        if body_started:
            return
        _add_section_break_new(doc, spec)
        body_section_index = len(doc.sections) - 1
        body_started = True
        last_was_page_break = True

    for section in export_ast.sections:
        if section.type == "cover":
            _add_cover_section(doc, section, spec, db=db)
            mark_content()
            if section.page_break_after:
                add_page_break_once()
        elif section.type == "toc":
            if section.page_break_before:
                add_page_break_once()
            _add_toc_section(doc, section, spec)
            mark_content()
            if section.page_break_after:
                add_page_break_once()
        elif section.type == "preface":
            if section.page_break_before:
                add_page_break_once()
            _add_publication_heading(doc, section.title, word_level=1, scale=scale)
            _render_blocks(doc, section.blocks, db=db, scale=scale)
            mark_content()
            if section.page_break_after:
                add_page_break_once()
        elif section.type == "chapter":
            first_body = not body_started
            begin_body_if_needed()
            if section.page_break_before and not first_body:
                add_page_break_once()
            _add_chapter_flyleaf(
                doc,
                section.title,
                getattr(section, "summary", "") or "",
                scale=scale,
            )
            mark_content()
            add_page_break_once()
            _render_blocks(doc, section.blocks, db=db, scale=scale, page_fresh=True)
            mark_content()
        elif section.type == "bibliography":
            begin_body_if_needed()
            if section.page_break_before:
                add_page_break_once()
            _add_publication_heading(doc, section.title, word_level=1, scale=scale)
            _render_blocks(doc, section.blocks, db=db, scale=scale)
            mark_content()

    configure_front_and_body_page_numbers(doc, body_section_index=body_section_index if body_started else 0)
    if not body_started:
        from app.services.publication.page_numbers import clear_section_footer

        for section in doc.sections:
            clear_section_footer(section)

    from app.services.publication.page_numbers import enable_update_fields_on_open

    enable_update_fields_on_open(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _render_blocks(
    doc: Document,
    blocks: list[AstBlock],
    db: Session | None = None,
    *,
    scale=None,
    page_fresh: bool = False,
) -> None:
    scale = scale or type_scale_for_format()
    # page_fresh：章过渡页后已分页，下一块从新页起
    fresh = page_fresh
    for block in blocks:
        role = block.role
        # 兼容旧数据：节不再独占页，按二级标题排
        if role == "section_flyleaf":
            role = "section_title"
        if role == "chapter_flyleaf":
            if not fresh:
                _add_page_break(doc)
            _add_chapter_flyleaf(
                doc,
                block.text,
                str(block.attrs.get("summary") or ""),
                scale=scale,
            )
            _add_page_break(doc)
            fresh = True
            continue

        if role == "body":
            node = block.attrs.get("tiptap_node")
            if isinstance(node, dict):
                docx_block(doc, node)
            else:
                _add_body(doc, block.text, scale=scale)
        elif role in ("section_title", "subsection_title"):
            # 每一节另起一页（标题下直接接正文，不独占空白页）
            if role == "section_title" and not fresh:
                _add_page_break(doc)
            node = block.attrs.get("tiptap_node")
            _add_publication_heading(
                doc,
                block.text,
                word_level=_word_heading_level_for_block(block) if block.role != "section_flyleaf" else 2,
                node=node if isinstance(node, dict) else None,
                scale=scale,
            )
        elif role == "figure":
            node = block.attrs.get("tiptap_node")
            if isinstance(node, dict):
                merged_attrs = merge_figure_export_attrs(block.attrs, node.get("attrs"))
                export_node = {**node, "attrs": merged_attrs}
                if not docx_figure_image_only(doc, export_node, db=db):
                    p = doc.add_paragraph("【图片待生成】")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph(block.text)
                for r in p.runs:
                    _style_docx_run(r, size_pt=scale.caption_pt, bold=True)
        elif role == "figure_caption":
            p = doc.add_paragraph(block.text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                _style_docx_run(r, size_pt=scale.caption_pt)
        elif role == "table_caption":
            p = doc.add_paragraph(block.text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                _style_docx_run(r, size_pt=scale.caption_pt, bold=True)
        elif role == "table" and block.attrs.get("table_node"):
            append_tiptap_to_document(doc, {"type": "doc", "content": [block.attrs["table_node"]]})
        elif role == "code":
            node = block.attrs.get("tiptap_node")
            if isinstance(node, dict):
                docx_block(doc, node)
            else:
                _add_code_paragraph(doc, block.text)
        elif role == "list":
            node = block.attrs.get("tiptap_node")
            if isinstance(node, dict):
                docx_block(doc, node)
        elif role == "blockquote":
            node = block.attrs.get("tiptap_node")
            if isinstance(node, dict):
                docx_block(doc, node)
            else:
                p = doc.add_paragraph(block.text)
                p.paragraph_format.left_indent = Pt(scale.first_indent_pt)
                p.paragraph_format.right_indent = Pt(scale.first_indent_pt)
                for r in p.runs:
                    _style_docx_run(r, size_pt=scale.body_pt, font_name=DOC_PREFACE_FONT)
        fresh = False


def render_ast_to_docx(ast: BookAst, db: Session | None = None) -> bytes:
    doc = Document()
    _init_docx_publication_fonts(doc)
    for block in ast.blocks:
        role = block.role
        if role == "book_title":
            p = doc.add_paragraph(block.text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                _style_docx_run(r, size_pt=BOOK_TITLE_PT, bold=True)
        elif role in ("preface_title", "chapter_title", "section_title", "subsection_title"):
            node = block.attrs.get("tiptap_node")
            _add_publication_heading(
                doc,
                block.text,
                word_level=_word_heading_level_for_block(block),
                node=node if isinstance(node, dict) else None,
            )
        elif role == "body":
            node = block.attrs.get("tiptap_node")
            if isinstance(node, dict):
                docx_block(doc, node)
            else:
                _add_body(doc, block.text)
        elif role == "figure":
            node = block.attrs.get("tiptap_node")
            if isinstance(node, dict):
                merged_attrs = merge_figure_export_attrs(block.attrs, node.get("attrs"))
                export_node = {**node, "attrs": merged_attrs}
                if not docx_figure_image_only(doc, export_node, db=db):
                    p = doc.add_paragraph("【图片待生成】")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p = doc.add_paragraph(block.text)
                for r in p.runs:
                    _style_docx_run(r, size_pt=CAPTION_PT, bold=True)
        elif role == "figure_caption":
            p = doc.add_paragraph(block.text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                _style_docx_run(r, size_pt=CAPTION_PT)
        elif role == "table_caption":
            p = doc.add_paragraph(block.text)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                _style_docx_run(r, size_pt=CAPTION_PT, bold=True)
        elif role == "table" and block.attrs.get("table_node"):
            append_tiptap_to_document(doc, {"type": "doc", "content": [block.attrs["table_node"]]})
        elif role == "code":
            node = block.attrs.get("tiptap_node")
            if isinstance(node, dict):
                docx_block(doc, node)
            else:
                _add_code_paragraph(doc, block.text)
        elif role == "list":
            node = block.attrs.get("tiptap_node")
            if isinstance(node, dict):
                docx_block(doc, node)
        elif role == "blockquote":
            node = block.attrs.get("tiptap_node")
            if isinstance(node, dict):
                docx_block(doc, node)
            else:
                p = doc.add_paragraph(block.text)
                p.paragraph_format.left_indent = Pt(18)
                for r in p.runs:
                    _style_docx_run(r, size_pt=BODY_PT)
    add_docx_page_numbers(doc)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
