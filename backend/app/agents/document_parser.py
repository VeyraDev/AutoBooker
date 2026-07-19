"""Parse PDF/DOCX/TXT, classify 资料 vs 参考文献；参考文献分块 embedding 入库。"""

from __future__ import annotations

import logging
import re
import uuid
from pathlib import Path
from typing import TYPE_CHECKING

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from sqlalchemy import select
from sqlalchemy.orm import Session

from app.config import settings
from app.llm.client import LLMClient
from app.models.book import Book
from app.models.reference import FileLifecycleStatus, ParseStatus, ReferenceChunk, ReferenceFile

if TYPE_CHECKING:
    pass

logger = logging.getLogger(__name__)

CHUNK_SIZE = 800
CHUNK_OVERLAP = 100

_MATERIAL_NAME_KW = ("要求", "约束", "风格", "说明", "不要", "注意", "术语")
_ACADEMIC_TEXT_KW = (
    "Abstract",
    "摘要",
    "关键词",
    "Keywords",
    "参考文献",
    "DOI",
    "Introduction",
    "Conclusion",
)


class DocumentParserAgent:
    def __init__(self, db: Session, book_id: uuid.UUID) -> None:
        self.db = db
        self.book_id = book_id
        self._client = LLMClient()

    @staticmethod
    def extract_text(file_path: str, file_type: str) -> str:
        ft = file_type.lower()
        if ft == "txt":
            return Path(file_path).read_text(encoding="utf-8", errors="replace")
        if ft == "pdf":
            doc = fitz.open(file_path)
            try:
                return "\n".join(page.get_text() for page in doc)
            finally:
                doc.close()
        if ft == "docx":
            from app.services.docx_math_extract import extract_docx_text_with_math

            try:
                return extract_docx_text_with_math(file_path)
            except Exception:
                logger.warning("docx math extract failed, fallback to plain text", exc_info=True)
                doc = DocxDocument(file_path)
                return "\n".join(p.text for p in doc.paragraphs if p.text and p.text.strip())
        raise ValueError(f"Unsupported file type: {file_type}")

    @staticmethod
    def classify_file(text: str, filename: str) -> str:
        """返回 material（资料型全文注入）或 reference（RAG）。"""
        word_count = len(text)
        lower_name = filename.lower()

        if any(kw in text for kw in _ACADEMIC_TEXT_KW):
            return "reference"
        if word_count < 5000 and any(kw in filename for kw in _MATERIAL_NAME_KW):
            return "material"
        if word_count < 5000:
            return "material"
        return "reference"

    @staticmethod
    def chunk_text(text: str) -> list[str]:
        if not text.strip():
            return []
        chunks: list[str] = []
        start = 0
        n = len(text)
        while start < n:
            end = min(start + CHUNK_SIZE, n)
            chunks.append(text[start:end])
            if end >= n:
                break
            start += CHUNK_SIZE - CHUNK_OVERLAP
        return chunks

    @classmethod
    def chunk_with_metadata(
        cls,
        file_path: str,
        file_type: str,
        fallback_text: str,
    ) -> list[dict]:
        """Split a source while preserving the best locator the format exposes."""
        ft = file_type.lower()
        rows: list[dict] = []
        if ft == "pdf":
            doc = fitz.open(file_path)
            try:
                for page_number, page in enumerate(doc, start=1):
                    for content in cls.chunk_text(page.get_text()):
                        rows.append({"content": content, "page_number": page_number})
            finally:
                doc.close()
            return rows
        if ft == "docx":
            doc = DocxDocument(file_path)
            heading_path: list[str] = []
            for paragraph_index, paragraph in enumerate(doc.paragraphs, start=1):
                text = paragraph.text.strip()
                if not text:
                    continue
                style = (paragraph.style.name if paragraph.style else "").lower()
                match = re.search(r"(?:heading|标题)\s*(\d+)", style)
                if match:
                    level = max(1, int(match.group(1)))
                    heading_path = heading_path[: level - 1] + [text]
                for content in cls.chunk_text(text):
                    rows.append(
                        {
                            "content": content,
                            "paragraph_index": paragraph_index,
                            "heading_path": list(heading_path),
                        }
                    )
            return rows
        heading_path: list[str] = []
        for paragraph_index, line in enumerate(fallback_text.splitlines(), start=1):
            text = line.strip()
            if not text:
                continue
            if re.match(r"^(?:#{1,6}\s+|第[一二三四五六七八九十百千\d]+章)", text):
                heading_path = [re.sub(r"^#{1,6}\s+", "", text)]
            for content in cls.chunk_text(text):
                rows.append(
                    {
                        "content": content,
                        "paragraph_index": paragraph_index,
                        "heading_path": list(heading_path),
                    }
                )
        return rows or [{"content": content} for content in cls.chunk_text(fallback_text)]

    def _embed_and_store_chunks(
        self,
        file_id: uuid.UUID,
        chunks: list[str],
        *,
        chunk_kind: str = "reference_material",
        metadata: list[dict] | None = None,
    ) -> None:
        embeddings: list[list[float]] = []
        batch_size = min(settings.EMBEDDING_BATCH_SIZE, 25)
        for i in range(0, len(chunks), batch_size):
            embeddings.extend(self._client.embed(chunks[i : i + batch_size]))

        if len(embeddings) != len(chunks):
            raise RuntimeError("embedding count mismatch")

        self.db.query(ReferenceChunk).filter(
            ReferenceChunk.file_id == file_id,
            ReferenceChunk.chunk_kind == chunk_kind,
        ).delete()

        for idx, (chunk, emb) in enumerate(zip(chunks, embeddings)):
            source = metadata[idx] if metadata and idx < len(metadata) else {}
            self.db.add(
                ReferenceChunk(
                    book_id=self.book_id,
                    file_id=file_id,
                    chunk_index=idx,
                    content=chunk,
                    embedding=emb,
                    chunk_kind=chunk_kind,
                    page_number=source.get("page_number"),
                    paragraph_index=source.get("paragraph_index", idx),
                    heading_path=source.get("heading_path"),
                    directly_quotable=chunk_kind == "citation_evidence",
                )
            )

    def parse_from_asset(
        self,
        file_id: uuid.UUID,
        asset_id: uuid.UUID,
        file_type: str,
        *,
        forced_class: str | None = None,
    ) -> None:
        from app.services.assets.binary_asset_service import BinaryAssetService
        from app.services.assets.temporary_workspace import TemporaryWorkspace

        asset = BinaryAssetService(self.db).get_asset_for_book(book_id=self.book_id, asset_id=asset_id)
        suffix = f".{file_type.lower()}"
        with TemporaryWorkspace().materialize(bytes(asset.content), suffix) as path:
            self.parse_and_store(file_id, str(path), file_type, forced_class=forced_class)

    def parse_and_store(
        self,
        file_id: uuid.UUID,
        file_path: str,
        file_type: str,
        *,
        forced_class: str | None = None,
    ) -> None:
        ref = self.db.get(ReferenceFile, file_id)
        if not ref or ref.book_id != self.book_id:
            logger.error("ReferenceFile %s not found for book %s", file_id, self.book_id)
            return

        ref.parse_status = ParseStatus.processing
        ref.lifecycle_status = FileLifecycleStatus.processing
        ref.error_message = None
        ref.parse_version = int(ref.parse_version or 0) + 1
        self.db.commit()

        try:
            text = self.extract_text(file_path, file_type)
            source_chunks = self.chunk_with_metadata(file_path, file_type, text)
            chunks = [str(item.get("content") or "") for item in source_chunks if item.get("content")]
            purposes = ref.file_purposes if isinstance(ref.file_purposes, list) else ["reference_material"]
            purposes = ["reference_material" if str(p) == "reference" else str(p) for p in purposes]

            if purposes:
                from app.services.material_parse_service import (
                    parse_file_artifacts,
                    persist_file_artifacts,
                    sync_material_conflicts,
                )

                artifacts = parse_file_artifacts(
                    text,
                    ref.filename or "",
                    [str(p) for p in purposes],
                    user_note=(ref.user_note or "").strip(),
                )
                artifacts["status"] = "effective"
                ref.parse_artifacts = artifacts
                persist_file_artifacts(self.db, ref, artifacts)

                if "reference_material" in purposes:
                    if chunks:
                        ref.ingest_kind = "reference"
                        self._embed_and_store_chunks(
                            file_id,
                            chunks,
                            chunk_kind="reference_material",
                            metadata=source_chunks,
                        )
                        artifacts["reference_chunk_count"] = len(chunks)
                    else:
                        ref.ingest_kind = "reference"
                else:
                    ref.ingest_kind = "material"

                if "bibliography" in purposes:
                    if chunks:
                        self._embed_and_store_chunks(
                            file_id,
                            chunks,
                            chunk_kind="citation_evidence",
                            metadata=source_chunks,
                        )

                from datetime import datetime, timezone

                ref.parse_status = ParseStatus.done
                ref.parsed_at = datetime.now(timezone.utc)
                ref.lifecycle_status = FileLifecycleStatus.effective
                self.db.commit()

                book = self.db.get(Book, self.book_id)
                if book:
                    conflicts = sync_material_conflicts(self.db, self.book_id)
                    artifacts["pending_issues"] = [
                        {
                            "type": item.get("type"),
                            "message": item.get("message"),
                        }
                        for item in conflicts
                        if str(ref.id) in (item.get("file_ids") or [])
                    ]
                    artifacts["status"] = "pending_confirmation" if artifacts["pending_issues"] else "effective"
                    ref.lifecycle_status = (
                        FileLifecycleStatus.pending_confirmation
                        if artifacts["pending_issues"]
                        else FileLifecycleStatus.effective
                    )
                    ref.parse_artifacts = artifacts
                    book.material_conflicts = conflicts if conflicts else book.material_conflicts
                    self.db.commit()

                if "bibliography" in purposes:
                    self._try_extract_bibliography_citations(text, file_id)
                return

            ref.ingest_kind = "reference"
            if not chunks:
                ref.parse_status = ParseStatus.done
                ref.error_message = "No text extracted"
                from datetime import datetime, timezone

                ref.parsed_at = datetime.now(timezone.utc)
                self.db.commit()
                return

            self._embed_and_store_chunks(file_id, chunks, metadata=source_chunks)

            from datetime import datetime, timezone

            ref.parse_status = ParseStatus.done
            ref.lifecycle_status = FileLifecycleStatus.effective
            ref.parsed_at = datetime.now(timezone.utc)
            self.db.commit()

            self._try_extract_bibliography_citations(text, file_id)
        except Exception as e:
            logger.exception("parse_and_store failed: %s", e)
            ref.parse_status = ParseStatus.failed
            ref.lifecycle_status = FileLifecycleStatus.failed
            ref.error_message = str(e)[:2000]
            self.db.commit()

    def retrieve(self, query: str, top_k: int = 5) -> list[str]:
        any_chunk = self.db.execute(
            select(ReferenceChunk.id).join(ReferenceFile, ReferenceChunk.file_id == ReferenceFile.id)
            .where(
                ReferenceChunk.book_id == self.book_id,
                ReferenceChunk.active.is_(True),
                ReferenceChunk.chunk_kind == "reference_material",
                ReferenceFile.parse_status == ParseStatus.done,
                ReferenceFile.lifecycle_status == FileLifecycleStatus.effective,
            )
            .limit(1)
        ).first()
        if not any_chunk:
            return []
        try:
            qvec = self._client.embed([query])[0]
        except Exception as e:
            logger.warning("embedding unavailable for RAG retrieve; skipping snippets: %s", e)
            return []
        stmt = (
            select(ReferenceChunk).join(ReferenceFile, ReferenceChunk.file_id == ReferenceFile.id)
            .where(
                ReferenceChunk.book_id == self.book_id,
                ReferenceChunk.active.is_(True),
                ReferenceChunk.chunk_kind == "reference_material",
                ReferenceFile.parse_status == ParseStatus.done,
                ReferenceFile.lifecycle_status == FileLifecycleStatus.effective,
            )
            .order_by(ReferenceChunk.embedding.cosine_distance(qvec))
            .limit(top_k)
        )
        rows = self.db.execute(stmt).scalars().all()
        return [r.content for r in rows]

    def retrieve_with_meta(self, query: str, top_k: int = 5) -> tuple[list[str], list[tuple[str, str]]]:
        """Returns (snippets, [(content, filename), ...])."""
        any_chunk = self.db.execute(
            select(ReferenceChunk.id).join(ReferenceFile, ReferenceChunk.file_id == ReferenceFile.id)
            .where(
                ReferenceChunk.book_id == self.book_id,
                ReferenceChunk.active.is_(True),
                ReferenceChunk.chunk_kind == "reference_material",
                ReferenceFile.parse_status == ParseStatus.done,
                ReferenceFile.lifecycle_status == FileLifecycleStatus.effective,
            )
            .limit(1)
        ).first()
        if not any_chunk:
            return [], []
        try:
            qvec = self._client.embed([query])[0]
        except Exception as e:
            logger.warning("embedding unavailable for RAG retrieve_with_meta; skipping snippets: %s", e)
            return [], []
        stmt = (
            select(ReferenceChunk, ReferenceFile.filename)
            .join(ReferenceFile, ReferenceChunk.file_id == ReferenceFile.id)
            .where(
                ReferenceChunk.book_id == self.book_id,
                ReferenceChunk.active.is_(True),
                ReferenceChunk.chunk_kind == "reference_material",
                ReferenceFile.parse_status == ParseStatus.done,
                ReferenceFile.lifecycle_status == FileLifecycleStatus.effective,
            )
            .order_by(ReferenceChunk.embedding.cosine_distance(qvec))
            .limit(top_k)
        )
        rows = self.db.execute(stmt).all()
        snippets: list[str] = []
        pairs: list[tuple[str, str]] = []
        for chunk_row, filename in rows:
            snippets.append(chunk_row.content)
            pairs.append((chunk_row.content, filename or ""))
        return snippets, pairs

    def _try_extract_bibliography_citations(self, text: str, file_id: uuid.UUID) -> None:
        """从参考文献型上传文件中解析书末条目并写入 citations 表。"""
        try:
            from app.agents.literature_agent import lookup_crossref_by_doi
            from app.models.book import Book
            from app.services.citation_service import ingest_uploaded_bibliography

            book = self.db.get(Book, self.book_id)
            if not book:
                return
            n = ingest_uploaded_bibliography(
                self.db,
                book,
                text,
                file_id,
                lookup_doi=lookup_crossref_by_doi,
            )
            if n:
                logger.info("extracted %s bibliography citations for book %s", n, self.book_id)
        except Exception as e:
            logger.warning("bibliography extraction skipped: %s", e)
