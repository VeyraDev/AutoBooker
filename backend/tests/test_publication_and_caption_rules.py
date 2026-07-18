from __future__ import annotations

from io import BytesIO
from types import SimpleNamespace
from uuid import uuid4

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml.ns import qn

from app.services import chapter_figure_table_normalize as normalizer
from app.services.figures.render.image_api.prompt_constraints import (
    build_direct_fallback_prompt,
    build_layout_agent_prompt,
)
from app.services.publication.book_ast import AstBlock, BookAst
from app.services.publication.publication_renderer_docx import render_ast_to_docx
from app.services.publication.publication_styles import PUBLICATION_CSS
from app.services.table_caption_ai import suggest_table_caption


def _heading_node(text: str, level: int) -> dict:
    return {
        "type": "heading",
        "attrs": {"level": level},
        "content": [{"type": "text", "text": text}],
    }


def _paragraph_outline_level(paragraph) -> str | None:
    ppr = paragraph._element.pPr
    if ppr is None:
        return None
    outline = ppr.find(qn("w:outlineLvl"))
    if outline is None:
        return None
    return outline.get(qn("w:val"))


def test_docx_export_heading_styles_follow_publication_hierarchy():
    ast = BookAst(
        title="书稿题名",
        blocks=[
            AstBlock(role="book_title", text="书稿题名"),
            AstBlock(role="preface_title", text="前言"),
            AstBlock(role="chapter_title", text="第一章 总论"),
            AstBlock(role="section_title", text="第一节", level=2, attrs={"tiptap_node": _heading_node("第一节", 2)}),
            AstBlock(role="subsection_title", text="一、", level=3, attrs={"tiptap_node": _heading_node("一、", 3)}),
            AstBlock(role="subsection_title", text="（一）", level=4, attrs={"tiptap_node": _heading_node("（一）", 4)}),
            AstBlock(role="subsection_title", text="1．", level=5, attrs={"tiptap_node": _heading_node("1．", 5)}),
            AstBlock(role="subsection_title", text="（1）", level=6, attrs={"tiptap_node": _heading_node("（1）", 6)}),
        ],
    )

    doc = Document(BytesIO(render_ast_to_docx(ast)))
    rows = [(p.text, p.alignment, p.runs[0].font.size.pt, p.runs[0].bold) for p in doc.paragraphs if p.text]

    assert rows[0] == ("书稿题名", WD_ALIGN_PARAGRAPH.CENTER, 22.0, True)
    assert rows[1] == ("前言", WD_ALIGN_PARAGRAPH.CENTER, 18.0, True)
    assert rows[2] == ("第一章 总论", WD_ALIGN_PARAGRAPH.CENTER, 18.0, True)
    assert rows[3] == ("第一节", WD_ALIGN_PARAGRAPH.CENTER, 16.0, True)
    assert rows[4] == ("一、", WD_ALIGN_PARAGRAPH.LEFT, 15.0, True)
    assert rows[5] == ("（一）", WD_ALIGN_PARAGRAPH.LEFT, 14.0, True)
    assert rows[6] == ("1．", WD_ALIGN_PARAGRAPH.LEFT, 12.0, False)
    assert rows[7] == ("（1）", WD_ALIGN_PARAGRAPH.LEFT, 12.0, False)

    book_title_p = doc.paragraphs[0]
    assert book_title_p.style.name == "Normal"
    assert _paragraph_outline_level(book_title_p) is None

    heading_expectations = [
        ("前言", "Heading 1", "0"),
        ("第一章 总论", "Heading 1", "0"),
        ("第一节", "Heading 2", "1"),
        ("一、", "Heading 3", "2"),
        ("（一）", "Heading 4", "3"),
        ("1．", "Heading 5", "4"),
        ("（1）", "Heading 6", "5"),
    ]
    by_text = {p.text: p for p in doc.paragraphs}
    for text, style_name, outline in heading_expectations:
        p = by_text[text]
        assert p.style.name == style_name
        assert _paragraph_outline_level(p) == outline


def test_pdf_export_css_matches_heading_hierarchy():
    assert "h1.chapter-title { font-size: 18pt; text-align: center; font-weight: bold" in PUBLICATION_CSS
    assert "h2.section-title { font-size: 16pt; text-align: center; font-weight: bold" in PUBLICATION_CSS
    assert "h3.subsection-title { font-size: 15pt; text-align: left; font-weight: bold" in PUBLICATION_CSS
    assert "h5.subsection-title { font-size: 12pt; text-align: left; font-weight: normal" in PUBLICATION_CSS
    assert "h6.subsection-title { font-size: 12pt; text-align: left; font-weight: normal" in PUBLICATION_CSS


def test_image_prompts_forbid_large_inner_titles():
    direct = build_direct_fallback_prompt("用户注册流程图：填写表单、邮件验证、完成注册", "process_flow")
    layout = build_layout_agent_prompt("用户注册流程图：填写表单、邮件验证、完成注册", "process_flow")

    assert "不要把“xxx图 / xxx示意图 / xxx流程图”作为画内标题" in direct
    assert "标题”写“无”" in layout


def _paragraph(text: str) -> dict:
    return {"type": "paragraph", "content": [{"type": "text", "text": text}]}


def _table() -> dict:
    def cell(text: str) -> dict:
        return {"type": "tableCell", "content": [_paragraph(text)]}

    return {
        "type": "table",
        "content": [
            {"type": "tableRow", "content": [cell("字段A"), cell("字段B")]},
            {"type": "tableRow", "content": [cell("10"), cell("20")]},
        ],
    }


class _Db:
    class _Query:
        def filter(self, *_args, **_kwargs):
            return self

        def first(self):
            return None

        def order_by(self, *_args, **_kwargs):
            return self

        def all(self):
            return []

    def query(self, *_args, **_kwargs):
        return self._Query()

    def add(self, _row):
        return None

    def flush(self):
        return None

    def commit(self):
        return None


def test_normalize_sort_uses_llm_titles_for_table_and_figure(monkeypatch):
    calls: list[str] = []
    monkeypatch.setattr(normalizer, "get_chapter_figures", lambda *_args, **_kwargs: [])
    monkeypatch.setattr(normalizer, "renumber_chapter_figures_from_tiptap", lambda *_args, **_kwargs: None)
    monkeypatch.setattr(normalizer, "suggest_table_caption", lambda *_args, **_kwargs: calls.append("table") or "核心指标表")
    monkeypatch.setattr(normalizer, "suggest_figure_caption", lambda *_args, **_kwargs: calls.append("figure") or "注册流程图")

    doc = {
        "type": "doc",
        "content": [
            _paragraph("本节比较两个关键指标。"),
            _table(),
            _paragraph("下面解释注册链路。"),
            {
                "type": "figureBlock",
                "attrs": {
                    "figureId": str(uuid4()),
                    "figureType": "flowchart",
                    "rawAnnotation": "请根据下面的布局脚本生成图片。用户注册流程图，包含填写表单、邮件验证、完成注册。",
                    "caption": "请根据下面的布局脚本生成图片。用户注册流程图，包含填写表单、邮件验证、完成注册。",
                },
            },
        ],
    }

    result = normalizer.normalize_chapter_figures_tables(
        uuid4(),
        1,
        doc,
        _Db(),
        book=SimpleNamespace(title="测试书稿"),
    )

    assert calls == ["table", "figure"]
    assert result["overview"][0]["title"] == "核心指标表"
    assert result["overview"][1]["title"] == "注册流程图"
    assert result["overview"][0]["title"] != "字段A、字段B"
    assert "请根据下面的布局脚本" not in result["overview"][1]["title"]


def test_normalize_sort_preview_mode_does_not_commit_or_persist(monkeypatch):
    class CountingDb(_Db):
        def __init__(self):
            self.commits = 0

        def commit(self):
            self.commits += 1

    def fail_persist(*_args, **_kwargs):
        raise AssertionError("preview mode must not persist figure changes")

    monkeypatch.setattr(normalizer, "get_chapter_figures", lambda *_args, **_kwargs: [])
    monkeypatch.setattr(normalizer, "ensure_figure_blocks_persisted", fail_persist)
    monkeypatch.setattr(normalizer, "renumber_chapter_figures_from_tiptap", fail_persist)
    monkeypatch.setattr(normalizer, "suggest_table_caption", lambda *_args, **_kwargs: "table title")
    monkeypatch.setattr(normalizer, "suggest_figure_caption", lambda *_args, **_kwargs: "figure title")

    doc = {
        "type": "doc",
        "content": [
            _paragraph("before table"),
            _table(),
            _paragraph("before figure"),
            {
                "type": "figureBlock",
                "attrs": {
                    "figureType": "flowchart",
                    "rawAnnotation": "figure flow",
                },
            },
        ],
    }
    db = CountingDb()

    result = normalizer.normalize_chapter_figures_tables(
        uuid4(),
        1,
        doc,
        db,
        book=SimpleNamespace(title="test book"),
        persist=False,
    )

    assert db.commits == 0
    assert result["overview"][0]["number"] == "1-1"
    assert result["overview"][1]["number"] == "1-1"
    assert "1-1" in result["text"]


def test_table_caption_fallback_never_uses_first_row(monkeypatch):
    class BrokenClient:
        def chat_completion(self, *_args, **_kwargs):
            raise RuntimeError("offline")

    monkeypatch.setattr("app.services.table_caption_ai.LLMClient", lambda: BrokenClient())

    title = suggest_table_caption(_table(), book=SimpleNamespace(title="测试书稿"), context="正文上下文")

    assert title == "本章数据表"
    assert title != "字段A、字段B"
