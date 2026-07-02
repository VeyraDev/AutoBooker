"""Tests for export page numbers."""

from io import BytesIO

import fitz
from docx import Document

from app.services.publication.book_ast import AstBlock, BookAst
from app.services.publication.page_numbers import add_docx_page_numbers, add_pdf_page_numbers
from app.services.publication.publication_renderer_docx import render_ast_to_docx
from app.services.publication.publication_renderer_pdf import render_ast_to_pdf


def _minimal_ast() -> BookAst:
    return BookAst(
        blocks=[
            AstBlock(role="book_title", text="测试书名"),
            AstBlock(role="body", text="第一段正文。"),
            AstBlock(role="body", text="第二段正文，用于分页测试。" * 40),
        ]
    )


def test_docx_export_has_page_number_field():
    doc = Document(BytesIO(render_ast_to_docx(_minimal_ast())))
    xml = doc._element.xml
    assert "fldChar" in xml
    assert " PAGE " in xml


def test_pdf_export_has_page_numbers():
    pdf_bytes = render_ast_to_pdf(_minimal_ast())
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    try:
        assert len(doc) >= 1
        text = doc[0].get_text()
        assert "1" in text
    finally:
        doc.close()


def test_add_pdf_page_numbers_on_existing_pdf():
    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "hello")
    raw = doc.tobytes()
    doc.close()
    out = add_pdf_page_numbers(raw)
    doc2 = fitz.open(stream=out, filetype="pdf")
    try:
        assert "1" in doc2[0].get_text()
    finally:
        doc2.close()


def test_add_docx_page_numbers():
    doc = Document()
    doc.add_paragraph("content")
    add_docx_page_numbers(doc)
    assert " PAGE " in doc._element.xml
