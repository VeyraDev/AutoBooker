"""DOCX 导出插图路径解析。"""

from __future__ import annotations

from docx import Document
from PIL import Image

from app.services.tiptap_convert import (
    _figure_raster_for_export,
    _resolve_figure_local_path,
    docx_figure_image_only,
    merge_figure_export_attrs,
)


def test_docx_figure_image_only_inserts_png(tmp_path):
    png = tmp_path / "figure.png"
    Image.new("RGB", (400, 300), "white").save(png)
    doc = Document()
    assert docx_figure_image_only(doc, {"type": "figureBlock", "attrs": {"fileUrl": str(png)}})
    assert doc.paragraphs[0].runs


def test_figure_raster_uses_png_sibling_for_svg(tmp_path):
    svg = tmp_path / "figure.svg"
    png = tmp_path / "figure.png"
    svg.write_text('<svg xmlns="http://www.w3.org/2000/svg" width="100" height="80"></svg>', encoding="utf-8")
    Image.new("RGB", (100, 80), "white").save(png)

    raster = _figure_raster_for_export(svg)
    assert raster is not None
    assert raster.suffix.lower() == ".png"

    doc = Document()
    assert docx_figure_image_only(doc, {"type": "figureBlock", "attrs": {"fileUrl": str(svg)}})


def test_resolve_static_url_strips_query(monkeypatch, tmp_path):
    from app.services.figures.storage import manager as storage_manager

    book_id = "c22dc3ec-a9b6-4895-9550-aa00fdbcc6c0"
    figure_id = "ea5c2f84-c4a8-47eb-a4d1-b5de5d08fafc"
    base = tmp_path / book_id / "1" / figure_id
    base.mkdir(parents=True)
    (base / "figure.svg").write_text(
        '<svg xmlns="http://www.w3.org/2000/svg" width="120" height="90"></svg>',
        encoding="utf-8",
    )

    class _Settings:
        @property
        def figures_path(self):
            return tmp_path

    monkeypatch.setattr(storage_manager, "settings", _Settings())
    url = f"/static/figures/{book_id}/1/{figure_id}/figure.svg?v=1780751319328"
    local = _resolve_figure_local_path({"figureId": figure_id, "fileUrl": url})
    assert local is not None
    assert local.suffix.lower() == ".svg"


def test_merge_figure_export_attrs_keeps_block_url_when_node_empty():
    block = {
        "figureId": "abc",
        "fileUrl": "/static/figures/book/1/abc/figure.svg",
        "tiptap_node": {"type": "figureBlock"},
    }
    node = {"fileUrl": "", "svgUrl": ""}
    merged = merge_figure_export_attrs(block, node)
    assert merged["fileUrl"] == block["fileUrl"]
    assert "tiptap_node" not in merged


def test_resolve_static_url_prefers_png(monkeypatch, tmp_path):
    from app.services.figures.storage import manager as storage_manager

    book_id = "c22dc3ec-a9b6-4895-9550-aa00fdbcc6c0"
    figure_id = "ea5c2f84-c4a8-47eb-a4d1-b5de5d08fafc"
    base = tmp_path / book_id / "1" / figure_id
    base.mkdir(parents=True)
    Image.new("RGB", (200, 150), "white").save(base / "figure.png")
    (base / "figure.svg").write_text("<svg></svg>", encoding="utf-8")

    class _Settings:
        @property
        def figures_path(self):
            return tmp_path

    monkeypatch.setattr(storage_manager, "settings", _Settings())

    attrs = {
        "figureId": figure_id,
        "fileUrl": f"/static/figures/{book_id}/1/{figure_id}/figure.svg",
    }
    local = _resolve_figure_local_path(attrs)
    assert local is not None
    assert local.name == "figure.png"
